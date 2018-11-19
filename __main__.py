# -*- coding: utf-8 -*-
"""
    main.py
    ~~~~~~~~~~~~~~
    :copyright:...
"""
# 导入需要的模块
import requests
import re
import io
import os
from bs4 import BeautifulSoup
from docx import Document
from docx import image
from docx.shared import Inches
from PIL import Image

path_seq = "\\" if os.name == 'nt' else "/"


class SpiderRenren(object):
    def __init__(self):
        # 登录人人网的url
        # url0 = "http://www.renren.com/PLogin.do"
        # 登录到个人主页的url
        self.headers = {
            'User-Agent': 'User-Agent: Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36'
        }
        url_3g = "http://3g.renren.com/login.do?autoLogin=true&&fx=0"
        data = {'email': "renggang@sina.com",
                'password': '---'}
        # 进行登录，并保存cookie
        self.req = requests.Session()
        login_res = self.req.post(url_3g, data=data, headers=self.headers)
        # 获取个人ID及信息
        login_bs = BeautifulSoup(login_res.text, "html.parser")
        home_elem = login_bs.select('div.sec.nav > a:nth-of-type(2)')[0]
        home_url = home_elem.get('href')
        # 访问个人主页
        home_res = self.req.get(home_url)
        self.home_bs = BeautifulSoup(home_res.text, "html.parser")

    def __del__(self):
        self.req.close()

    def get_photo(self):
        # 进入照片集列表页
        photos_home_elem = self.home_bs.select('table')[2].select('tr')[0].select('td')[1]
        photos_home_url = photos_home_elem.select('a')[0].get('href')
        photos_home_num = photos_home_elem.select('span')[0].getText()
        photos_home_res = self.req.get(photos_home_url, headers=self.headers)
        photos_home_bs = BeautifulSoup(photos_home_res.text, "html.parser")
        # 把照片集全部找出来，最后一个div是翻页
        is_photos_next = True
        photos_current_bs = photos_home_bs
        while is_photos_next:
            photos_elem_list = photos_current_bs.select('div.list')[0].select('div')
            # 遍历照片集
            for photos_elem in photos_elem_list[1:-1]:
                photos_url = photos_elem.select('a')[-1].get('href')
                photos_title = photos_elem.select('a')[-1].getText().strip()
                photos_update = re.split("\s+", photos_elem.select('span.ns')[0].getText().strip())[1]
                path = photos_title + photos_update
                if not os.path.exists(path):
                    os.mkdir(path)
                photos_res = self.req.get(photos_url, headers=self.headers)
                photos_bs = BeautifulSoup(photos_res.text, "html.parser")
                # 进入单张浏览页面
                try:
                    photo_list_url = photos_bs.select('div.list')[0].select('a')[0].get('href')
                except IndexError:
                    # 此照片集内无照片
                    continue
                photo_list_res = self.req.get(photo_list_url, headers=self.headers)
                photo_list_bs = BeautifulSoup(photo_list_res.text, "html.parser")
                # 判断是否有下一张，如果有，会继续循环
                is_photo_next = True
                photo_current_bs = photo_list_bs
                id = 0
                while is_photo_next:
                    photo_title = \
                        re.split("\s+", photo_current_bs.select('div.sec')[2].select('p')[1].getText().strip())[0]
                    # 有可能出现没有标题的图片，用自增的id表示
                    if photo_title == "小图":
                        photo_title = str(id)
                        id += 1
                    photo_date_str_list = re.split("\s+|:",
                                                   photo_current_bs.select('div.sec')[2].select('p')[
                                                       2].getText().strip())
                    photo_date = photo_date_str_list[1] + photo_date_str_list[2] + ":" + photo_date_str_list[3]
                    photo_location = ""
                    document = Document()
                    document.add_heading(photo_title, 0)
                    document.add_paragraph(photo_date)
                    img_url = photo_current_bs.select('div.sec')[2].select('p')[1].select('a')[-1].get('href')
                    photo_content = self._get_img_via_url(img_url)
                    try:
                        document.add_picture(photo_content, width=Inches(6))
                    except image.exceptions.UnrecognizedImageError:
                        # img无法下载，forbidden403
                        pass
                    document.add_heading('好友评论', level=1)
                    # 获取评论
                    try:
                        photo_comment_list = photo_current_bs.select('div.list')[0].select('div')
                        self._process_comments(document, photo_comment_list)
                    except IndexError:
                        # no comments
                        pass
                    document.add_page_break()
                    document.save(path + path_seq + photo_title + '.docx')
                    # 获取下一页信息
                    try:
                        photo_is_next_list = re.findall("\d+", photo_current_bs.select('div.sec')[3].select('span')[
                            -1].getText().strip())
                    except IndexError:
                        # 只有1页的情况下，没有下一页的标签
                        is_photo_next = False
                        continue
                    is_photo_next = True if photo_is_next_list[0] != photo_is_next_list[1] else False
                    if is_photo_next:
                        photo_next_url = photo_current_bs.select('div.sec')[3].a.get("href")
                        photo_next_res = self.req.get(photo_next_url, headers=self.headers)
                        photo_next_bs = BeautifulSoup(photo_next_res.text, "html.parser")
                        photo_current_bs = photo_next_bs
            is_photos_next = True if photos_elem_list[-1].select('a')[0].get('title') == "下一页" else False
            if is_photos_next:
                photos_next_url = photos_elem_list[-1].select('a')[0].get('href')
                photos_next_res = self.req.get(photos_next_url, headers=self.headers)
                photos_next_bs = BeautifulSoup(photos_next_res.text, "html.parser")
                photos_current_bs = photos_next_bs

    def get_blog(self):
        # 获取日志
        blog_home_elem = self.home_bs.select('table')[2].select('tr')[0].select('td')[0]
        blog_home_url = blog_home_elem.select('a')[0].get('href')
        blog_home_num = blog_home_elem.select('span')[0].getText()
        blog_home_res = self.req.get(blog_home_url, headers=self.headers)
        blog_home_bs = BeautifulSoup(blog_home_res.text, "html.parser")
        # 把日志列表全部找出来，最后一个div是翻页
        is_next = True
        blog_current_bs = blog_home_bs
        while is_next:
            blog_elem_list = blog_current_bs.select('div.list')[0].select('div')
            for blog_elem in blog_elem_list[:-1]:
                blog_url = blog_elem.select('a')[0].get('href')
                blog_res = self.req.get(blog_url, headers=self.headers)
                blog_bs = BeautifulSoup(blog_res.text, "html.parser")
                blog_title = blog_bs.select('div.note')[0].b.getText().strip()
                blog_date = "".join(re.split("\s+", blog_bs.select('div.note')[0].p.getText().strip())[0:2])
                document = Document()
                document.add_heading(blog_title, 0)
                document.add_paragraph(blog_date)
                document.add_heading('正文', level=1)
                # 都是</br>，无法用children选出blog内容，尝试正则的方式
                blog_content_list = blog_bs.select('div.con')[0].contents
                for blog_content in blog_content_list:
                    if blog_content.name == "img":
                        img_url = blog_content.get('src')
                        img_res = self.req.get(img_url, headers=self.headers)
                        img_data = img_res.content
                        img_io = io.BytesIO()
                        img_io.write(img_data)
                        img_io.seek(0)
                        im = Image.open(img_io)  # 通过im.size可以获取图片的尺寸，此处我没有用到
                        # 1024像素=3.413英寸=8.67厘米
                        document.add_picture(img_io, width=Inches(3))
                    elif blog_content.name == "br":
                        pass
                    else:
                        document.add_paragraph(blog_content)
                document.add_heading('好友评论', level=1)
                # 获取评论
                try:
                    blog_comment_list = blog_bs.select('div.list')[0].select('div')
                    self._process_comments(document, blog_comment_list)
                except IndexError:
                    # no comments
                    pass
                document.add_page_break()
                document.save(blog_title + '.docx')
            is_next = True if blog_elem_list[-1].select('a')[0].get('title') == "下一页" else False
            if is_next:
                blog_next_url = blog_elem_list[-1].select('a')[0].get('href')
                blog_next_res = self.req.get(blog_next_url, headers=self.headers)
                blog_next_bs = BeautifulSoup(blog_next_res.text, "html.parser")
                blog_current_bs = blog_next_bs

    def get_share(self):
        # 获取分享
        share_elem = self.home_bs.select('table')[2].select('tr')[1].select('td')[0]
        share_url = share_elem.select('a')[0].get('href')
        share_num = share_elem.select('span')[0].getText()

    def get_status(self):
        # 获取状态
        status_elem = self.home_bs.select('table')[2].select('tr')[1].select('td')[1]
        status_url = status_elem.select('a')[0].get('href')
        status_num = status_elem.select('span')[0].getText()

    """
       在线获取图片并返回
    """

    def _get_img_via_url(self, url):
        img_res = self.req.get(url, headers=self.headers)
        img_data = img_res.content
        img_io = io.BytesIO()
        img_io.write(img_data)
        img_io.seek(0)
        return img_io

    def _process_comments(self, document, comment_list):
        for comment_div in comment_list:
            for comment in comment_div.contents:
                if comment.name == "a":
                    document.add_paragraph(comment.getText())
                elif comment.name == "br":
                    continue
                elif comment.name == "p":
                    document.add_paragraph(' '.join(re.split('\s+', comment.getText().strip())[0:2]))
                elif comment.name == "img":
                    # 表情
                    img_url = comment.get('src')
                    photo_content = self._get_img_via_url(img_url)
                    document.add_picture(photo_content, width=Inches(0.2))
                else:
                    if comment.strip():
                        document.add_paragraph(comment.strip())


if __name__ == '__main__':
    s = SpiderRenren()
    s.get_blog()
